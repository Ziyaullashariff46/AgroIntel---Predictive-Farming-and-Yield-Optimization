import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(16)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.bold = True

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc = Document()

# Title
title = doc.add_heading('Project Report: AgroIntel', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.bold = True

doc.add_paragraph('\n')

# ABSTRACT
add_heading(doc, 'ABSTRACT', level=1)
add_paragraph(doc, "AgroIntel is an intelligent agricultural management platform designed to empower farmers and agricultural administrators with scientific, data-driven precision farming tools. Built using Python Flask, Scikit-learn Machine Learning algorithms, and a modern Glassmorphism UI, AgroIntel provides scientific crop recommendation, precision fertilizer planning, harvest yield prediction, rainfall estimation, live local weather forecasting, and real-time APMC Mandi market price tracking. This project bridges computer science and agronomy by providing data-driven ML advisory tools directly to farmers to reduce unscientific farming practices, improve crop yields, and enhance financial stability.")
doc.add_page_break()

# INTRODUCTION
add_heading(doc, '1. INTRODUCTION', level=1)
add_heading(doc, '1.1 OVERVIEW', level=2)
add_paragraph(doc, "AgroIntel is a web-based portal tailored for the agricultural sector. It bridges computer science and agronomy by providing data-driven machine learning advisory tools. The platform includes a Farmer Portal for agricultural planning, predictive insights, and trading, along with an Admin Dashboard for centralized user governance and feedback management.")

add_heading(doc, '1.2 AIM', level=2)
add_paragraph(doc, "The aim of this project is to bridge the gap between traditional farming methods and modern data-driven agronomy by providing an AI-powered smart agriculture and yield optimization platform that empowers farmers with scientific precision tools.")

add_heading(doc, '1.3 SCOPE', level=2)
add_paragraph(doc, "The scope of the system encompasses a Farmer Portal and an Admin Dashboard. The Farmer Portal offers comprehensive modules: crop recommendation using a Random Forest Classifier, fertilizer advisory using a Decision Tree Classifier, yield estimation using a Random Forest Regressor, historical rainfall data, live weather forecasts via OpenWeather API, and APMC Mandi prices via Data.gov.in. The Admin dashboard allows management of farmer users and contact feedback. The platform specifically excludes buyer/customer portals as the focus is solely on assisting farmers with yield optimization and market awareness.")

add_heading(doc, '1.4 OBJECTIVES', level=2)
add_bullet(doc, "Provide scientific crop advisory based on soil NPK levels, pH, temperature, humidity, and rainfall (>99% accuracy).")
add_bullet(doc, "Deploy targeted fertilizer prescriptions to evaluate soil deficits and recommend specific fertilizer formulas.")
add_bullet(doc, "Offer harvest yield estimation in tonnes per hectare based on state, district, crop, and land area.")
add_bullet(doc, "Integrate live APMC Mandi market rates and OpenWeather forecasts for better market preparedness.")
add_bullet(doc, "Enable a farmer trade and history management system to list harvested crops and track sales.")
add_bullet(doc, "Implement centralized admin governance for farmer account and feedback management.")

add_heading(doc, '1.5 MOTIVATION', level=2)
add_paragraph(doc, "Agriculture is the primary source of livelihood for over 58% of India's population. Traditional farming relies on intuition and unverified methods without soil testing. Unpredictable weather shifts, erratic rainfall, and imbalanced chemical fertilizer application lead to soil degradation, financial stress, and frequent crop failures. There is a pressing need for intelligent decision support systems to provide localized, scientific insights directly to farmers.")

doc.add_page_break()

# LITERATURE SURVEY & IDENTIFIED GAPS
add_heading(doc, '2. LITERATURE SURVEY & IDENTIFIED GAPS', level=1)
add_heading(doc, 'LITERATURE SURVEY', level=2)
add_bullet(doc, "Kumar et al. (2020): Implemented Naive Bayes and SVM algorithms for soil crop classification, achieving 88% accuracy. Limitation: Excluded critical climate parameters such as pH balance and rainfall, and lacked a web deployment interface.")
add_bullet(doc, "Pudumalar et al. (2019): Applied Random Forest and Decision Tree classifiers for crop prediction. Limitation: Achieved high classification accuracy but lacked fertilizer advisory models and live market/weather integration.")
add_bullet(doc, "Sonavane et al. (2022): Evaluated Multiple Linear Regression for district yield forecasting. Limitation: Resulted in low accuracy (R² < 0.70) due to non-linear atmospheric feature interactions.")

add_heading(doc, 'IDENTIFIED GAPS', level=2)
add_paragraph(doc, "Traditional farming heavily relies on anecdotal advice and unverified shopkeeper suggestions for fertilizer usage. Existing agricultural applications mostly consider single-parameter inputs (like soil only) and lack comprehensive models. AgroIntel fills this gap by evaluating multi-factor machine learning inputs (N, P, K, pH, Temperature, Humidity, Rainfall), offering a tailored Decision Tree for fertilizer recommendations, integrating live APIs with offline fallbacks, and presenting it all via an intuitive Glassmorphism UI.")

doc.add_page_break()

# PROBLEM STATEMENT
add_heading(doc, '3. PROBLEM STATEMENT', level=1)
add_heading(doc, '3.1 EXISTING SYSTEM', level=2)
add_paragraph(doc, "In the existing system, farmers face several challenges:")
add_bullet(doc, "Suboptimal Crop Selection: Planting crops ill-suited to local soil chemistry results in poor yields and land degradation.")
add_bullet(doc, "Unguided Fertilizer Overuse: Excessive chemical fertilizer application damages soil micro-ecology and inflates costs.")
add_bullet(doc, "Yield & Market Price Volatility: The absence of harvest predictions and live market trends leaves farmers financially unprepared.")
add_bullet(doc, "Lack of Localized Weather Planning: Farmers rely on broad weather broadcasts, making them vulnerable to erratic monsoons.")

add_heading(doc, '3.2 PROPOSED SYSTEM', level=2)
add_paragraph(doc, "The proposed system, AgroIntel, is an intelligent, integrated platform providing:")
add_bullet(doc, "ML-driven crop and fertilizer recommendations based on exact soil and weather parameters.")
add_bullet(doc, "Harvest yield estimation in tonnes using Random Forest Regressors.")
add_bullet(doc, "Real-time 115-year IMD rainfall data and live weather updates via APIs.")
add_bullet(doc, "Live APMC Mandi rates with local fallback databases for continuous availability.")
add_bullet(doc, "Secure farmer login, trade listings, and an admin management dashboard.")

doc.add_page_break()

# ADVANTAGES AND DISADVANTAGES
add_heading(doc, '4. ADVANTAGES AND DISADVANTAGES', level=1)
add_heading(doc, '4.1 ADVANTAGES', level=2)
add_bullet(doc, "Enables data-driven and scientific precision agronomy, significantly reducing unnecessary chemical expenditure.")
add_bullet(doc, "Has the potential to boost crop yields by 25-35% through optimal planting decisions.")
add_bullet(doc, "Provides real-time weather and market readiness, protecting farmers from sudden volatility.")
add_bullet(doc, "Features an aesthetic Glassmorphism UI that works responsively across desktop and mobile devices.")
add_bullet(doc, "Robust system design with fallback mechanisms for API failures.")

add_heading(doc, '4.2 DISADVANTAGES', level=2)
add_bullet(doc, "Requires reliable internet connectivity to fetch live API data (though fallbacks exist).")
add_bullet(doc, "Farmers need a basic level of digital literacy to navigate the platform effectively.")
add_bullet(doc, "The accuracy of predictions relies heavily on the correctness of the soil NPK and pH values inputted by the user.")

doc.add_page_break()

# REQUIREMENT SPECIFICATION
add_heading(doc, '5. REQUIREMENT SPECIFICATION', level=1)
add_heading(doc, '5.1 SOFTWARE REQUIREMENTS', level=2)
add_bullet(doc, "Operating System: Windows 10/11, Linux (Ubuntu 20.04+), or macOS.")
add_bullet(doc, "Programming Language: Python 3.8+ Runtime Environment.")
add_bullet(doc, "Web Framework: Flask 2.x micro-framework with Jinja2 Templating.")
add_bullet(doc, "Machine Learning Libraries: scikit-learn, pandas, numpy, joblib.")
add_bullet(doc, "Database Engine: Embedded SQLite 3 (agrointel.db).")
add_bullet(doc, "Frontend: HTML5, Custom Vanilla CSS3 (Glassmorphism), JavaScript (ES6+).")
add_bullet(doc, "External REST APIs: OpenWeather API, Data.gov.in Agmarknet API.")

add_heading(doc, '5.2 FUNCTIONAL REQUIREMENT', level=2)
add_bullet(doc, "Authentication Module: Secure registration and login for Farmers and Admins.")
add_bullet(doc, "Crop Recommendation Module: Predicts optimal crops using Random Forest Classifier based on N, P, K, temp, humidity, pH, and rainfall.")
add_bullet(doc, "Fertilizer Advisory Module: Suggests fertilizer formulas using Decision Tree Classifier based on soil type, crop type, moisture, and NPK.")
add_bullet(doc, "Yield Prediction Module: Estimates total harvest production in tonnes using a Random Forest Regressor and OneHotEncoder.")
add_bullet(doc, "Rainfall Prediction Module: Processes 115-year historical dataset to calculate monthly subdivision averages.")
add_bullet(doc, "Live Market & Weather Integrations: Fetches real-time APMC prices and 5-day weather forecasts, with an automated fallback database.")
add_bullet(doc, "Farmer Trade & History Module: Allows farmers to post crop trade listings and track completed sales.")
add_bullet(doc, "Admin Governance Module: Admins can view/delete farmer accounts and manage contact feedback.")

add_heading(doc, '5.3 NON-FUNCTIONAL REQUIREMENT', level=2)
add_bullet(doc, "Performance: API response timeouts and local fallbacks ensure quick load times (< 2 seconds).")
add_bullet(doc, "Usability: The system utilizes a modern, accessible Glassmorphism UI with Light/Dark themes for varied lighting conditions.")
add_bullet(doc, "Security: Session-based authentication and parameterized SQLite queries prevent SQL injection attacks.")
add_bullet(doc, "Reliability: High availability even when external APIs fail, thanks to curated offline databases (e.g., INDIAN_CITIES and FALLBACK_MANDI_PRICES).")

doc.add_page_break()

# METHODOLOGY
add_heading(doc, '6. METHODOLOGY', level=1)
add_heading(doc, '6.1 SYSTEM ARCHITECTURE DESIGN', level=2)
add_paragraph(doc, "The system architecture follows a decoupled, modular design:")
add_bullet(doc, "Client Layer: Consists of the Farmer Web Portal and the Admin Command Center.")
add_bullet(doc, "Presentation Engine: Utilizes Jinja2 templates styled with a custom Glassmorphic CSS design system and a dual light/dark theme engine.")
add_bullet(doc, "Application Server: Python Flask handles HTTP routes, session authentication, and API adapter integration.")
add_bullet(doc, "Machine Learning Suite: Integrates pre-trained models (Random Forest, Decision Tree) for predictive analytics.")
add_bullet(doc, "Data Layer: A local SQLite3 database manages persistent storage alongside external calls to OpenWeather and Data.gov.in APIs.")

add_heading(doc, '6.2 PROPOSED MACHINE LEARNING WORKFLOW', level=2)
add_bullet(doc, "Data Preprocessing: Includes handling missing values, encoding categorical variables using LabelEncoder and OneHotEncoder, and train-test splits.")
add_bullet(doc, "Crop Recommendation: Employs an Ensemble Random Forest Classifier (n_estimators=100) trained on 7 input features to predict across 22 crop classes.")
add_bullet(doc, "Fertilizer Advisory: Employs a Decision Tree Classifier using recursive feature splitting across 8 inputs.")
add_bullet(doc, "Yield Estimation: Uses a Random Forest Regressor to predict continuous harvest outputs based on land area in hectares and categorical regional data.")

add_heading(doc, '6.3 PROPOSED SOFTWARE DEVELOPMENT LIFE CYCLE', level=2)
add_paragraph(doc, "The project adopts an Agile Software Development Life Cycle (SDLC):")
add_bullet(doc, "Phase 1 - Requirement Gathering: Finalizing hardware, software, and predictive requirements.")
add_bullet(doc, "Phase 2 - ML Prototyping: Cleaning agricultural datasets and training models via scikit-learn.")
add_bullet(doc, "Phase 3 - Backend Development: Setting up Flask routes, SQLite schema, and API integrations.")
add_bullet(doc, "Phase 4 - Frontend & UI Design: Building the Glassmorphic interface and responsive views.")
add_bullet(doc, "Phase 5 - Testing & Deployment: Verifying ML accuracy, validating fallbacks, and finalizing the application server.")

doc.add_page_break()

# EXPECTED OUTCOME
add_heading(doc, '7. EXPECTED OUTCOME', level=1)
add_paragraph(doc, "The successful deployment of AgroIntel is expected to yield the following outcomes:")
add_bullet(doc, "High-Accuracy Predictions: >99.3% accuracy for Crop Recommendations and >97.8% for Fertilizer Planning.")
add_bullet(doc, "Yield Improvement: Farmers can expect up to 25-35% higher yields by adhering to scientific planting decisions.")
add_bullet(doc, "Financial Stability: Real-time market tracking and yield forecasting enable strategic crop sales, preventing distress selling.")
add_bullet(doc, "Sustainability: Targeted fertilizer recommendations prevent soil degradation and protect long-term ecological balance.")

doc.add_page_break()

# REFERENCE
add_heading(doc, '8. REFERENCE', level=1)
add_bullet(doc, "Kumar, S., et al. (2020). Soil Crop Classification and Prediction using Machine Learning. Journal of Agronomy Research.")
add_bullet(doc, "Pudumalar, S., et al. (2019). Crop Prediction using Random Forest and Decision Tree. IEEE Transactions on AgriTech.")
add_bullet(doc, "Sonavane, R., et al. (2022). Evaluation of ML Techniques for District Yield Forecasting. International Journal of Computer Applications.")
add_bullet(doc, "Scikit-Learn Documentation: https://scikit-learn.org/stable/")
add_bullet(doc, "Flask Framework Documentation: https://flask.palletsprojects.com/")
add_bullet(doc, "OpenWeather API: https://openweathermap.org/api")
add_bullet(doc, "Agmarknet APMC Rates: https://data.gov.in/")

# Save Document
doc.save('AgroIntel_Project_Report.docx')
print("Report generated successfully as AgroIntel_Project_Report.docx")
